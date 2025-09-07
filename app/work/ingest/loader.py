# app/rag_app/ingest/loader.py
# -*- coding: utf-8 -*-
from __future__ import annotations
import os, io, hashlib, mimetypes
from typing import Dict, Any, List, Optional

#try:
from pdf_parser import pdf_parser as pdf_parser_raw
#except Exception:
    #pdf_parser_raw = None

try:
    import docx  # python-docx
except Exception:
    docx = None

# ---------- единый целевой формат ----------
# ParsedDoc:
# {
#   "doc_id": str,
#   "source_path": str,
#   "mime": str,
#   "title": Optional[str],
#   "meta": dict,
#   "pages": [ {"page": int, "text": str, "tables_md": list[str]} ]
# }

# ---------- утилиты ----------

def _file_sha256(path: str, chunk_size: int = 1 << 20) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(chunk_size), b""):
            h.update(chunk)
    return h.hexdigest()

def _guess_mime(path: str) -> str:
    mime, _ = mimetypes.guess_type(path)
    return mime or "application/octet-stream"

def _norm_title_from_filename(path: str) -> str:
    name = os.path.basename(path)
    return os.path.splitext(name)[0]

# ---------- адаптер для pdf_parser ----------

def _parse_pdf_to_parsed_doc(path: str) -> Dict[str, Any]:
    if pdf_parser_raw is None:
        raise ImportError("Модуль 'pdf_parser' не найден. Убедись, что pdf_parser.py в PYTHONPATH.")

    pages_map = pdf_parser_raw(path, image_flag=False)
    pages_out: List[Dict[str, Any]] = []
    keys = sorted(pages_map.keys(), key=lambda k: int(k.split("_")[-1]))
    for k in keys:
        page_idx = int(k.split("_")[-1]) + 1
        page_text_blocks = pages_map[k][0] if len(pages_map[k]) > 0 else []
        page_tables_strs = pages_map[k][3] if len(pages_map[k]) > 3 else []
        text = "".join(page_text_blocks).strip()
        pages_out.append({
            "page": page_idx,
            "text": text,
            "tables_md": page_tables_strs or [],
        })

    return {
        "doc_id": _file_sha256(path),
        "source_path": os.path.abspath(path),
        "mime": "application/pdf",
        "title": _norm_title_from_filename(path),
        "meta": {"num_pages": len(pages_out)},
        "pages": pages_out,
    }

# ---------- загрузчики для других типов ----------

def _parse_txt_to_parsed_doc(path: str) -> Dict[str, Any]:
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    pages = [{"page": 1, "text": text, "tables_md": []}]
    return {
        "doc_id": _file_sha256(path),
        "source_path": os.path.abspath(path),
        "mime": "text/plain",
        "title": _norm_title_from_filename(path),
        "meta": {"num_pages": 1},
        "pages": pages,
    }

def _parse_md_to_parsed_doc(path: str) -> Dict[str, Any]:
    return _parse_txt_to_parsed_doc(path) | {"mime": "text/markdown"}

def _parse_html_to_parsed_doc(path: str) -> Dict[str, Any]:
    try:
        from bs4 import BeautifulSoup
    except Exception:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return {
            "doc_id": _file_sha256(path),
            "source_path": os.path.abspath(path),
            "mime": "text/html",
            "title": _norm_title_from_filename(path),
            "meta": {"num_pages": 1},
            "pages": [{"page": 1, "text": text, "tables_md": []}],
        }
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    text = soup.get_text("\n")
    return {
        "doc_id": _file_sha256(path),
        "source_path": os.path.abspath(path),
        "mime": "text/html",
        "title": soup.title.string.strip() if soup.title and soup.title.string else _norm_title_from_filename(path),
        "meta": {"num_pages": 1},
        "pages": [{"page": 1, "text": text, "tables_md": []}],
    }

def _parse_docx_to_parsed_doc(path: str) -> Dict[str, Any]:
    if docx is None:
        raise ImportError("Пакет 'python-docx' не установлен.")
    d = docx.Document(path)
    paras = [p.text for p in d.paragraphs]
    text = "\n".join(paras).strip()
    return {
        "doc_id": _file_sha256(path),
        "source_path": os.path.abspath(path),
        "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "title": _norm_title_from_filename(path),
        "meta": {"num_pages": 1},
        "pages": [{"page": 1, "text": text, "tables_md": []}],
    }

# ---------- главная функция ----------

def extract_text(path: str) -> Dict[str, Any]:
    if not os.path.exists(path):
        raise FileNotFoundError(path)

    ext = os.path.splitext(path.lower())[1]
    if ext in [".pdf"]:
        return _parse_pdf_to_parsed_doc(path)
    elif ext in [".txt"]:
        return _parse_txt_to_parsed_doc(path)
    elif ext in [".md", ".markdown"]:
        return _parse_md_to_parsed_doc(path)
    elif ext in [".html", ".htm"]:
        return _parse_html_to_parsed_doc(path)
    elif ext in [".docx"]:
        return _parse_docx_to_parsed_doc(path)
    else:
        mime = _guess_mime(path)
        if mime == "application/pdf":
            return _parse_pdf_to_parsed_doc(path)
        elif mime.startswith("text/"):
            return _parse_txt_to_parsed_doc(path) | {"mime": mime}
        raise ValueError(f"Неподдерживаемый тип файла: {ext or mime}")

# ---------- адаптер в LlamaIndex ----------

def to_llama_documents(parsed_doc: Dict[str, Any]):
    try:
        from llama_index.core import Document
    except Exception:
        docs = []
        for p in parsed_doc["pages"]:
            docs.append({
                "text": p["text"],
                "metadata": {
                    "doc_id": parsed_doc["doc_id"],
                    "page": p["page"],
                    "title": parsed_doc.get("title"),
                    "source": parsed_doc["source_path"],
                    "mime": parsed_doc["mime"],
                    "tables_md": p.get("tables_md", []),
                }
            })
        return docs

    out = []
    for p in parsed_doc["pages"]:
        out.append(
            Document(
                text=p["text"],
                metadata={
                    "doc_id": parsed_doc["doc_id"],
                    "page": p["page"],
                    "title": parsed_doc.get("title"),
                    "source": parsed_doc["source_path"],
                    "mime": parsed_doc["mime"],
                    "tables_md": p.get("tables_md", []),
                }
            )
        )
    return out
